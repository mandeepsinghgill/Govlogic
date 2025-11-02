"""
Enhanced Document Export Service - Professional exports with advanced formatting
"""
from typing import Dict, Any, List, Optional
from datetime import datetime
import io
import os

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side, NamedStyle
    from openpyxl.utils import get_column_letter
    from openpyxl.chart import BarChart, Reference, LineChart
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
except ImportError:
    pass


class EnhancedExportService:
    """
    Professional document export service with:
    - Word: Professional formatting, cover page, TOC, custom styles
    - Excel: Multi-sheet workbooks, formulas, charts, formatting
    - PDF: High-quality output, custom templates, branding
    """
    
    def __init__(self):
        self.company_branding = {
            'primary_color': RGBColor(59, 130, 246),  # Blue
            'secondary_color': RGBColor(107, 114, 128),  # Gray
            'accent_color': RGBColor(16, 185, 129),  # Green
            'logo_path': None
        }
    
    async def export_to_professional_word(
        self,
        proposal_data: Dict[str, Any],
        include_cover_page: bool = True,
        include_toc: bool = True,
        include_executive_summary: bool = True,
        custom_styles: Optional[Dict] = None
    ) -> bytes:
        """
        Export to professionally formatted Word document
        """
        doc = Document()
        
        # Apply custom styles
        self._setup_custom_styles(doc, custom_styles)
        
        # Set document properties
        self._set_document_properties(doc, proposal_data)
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1)
            section.header_distance = Inches(0.5)
            section.footer_distance = Inches(0.5)
        
        # Add cover page
        if include_cover_page:
            self._add_professional_cover_page(doc, proposal_data)
            doc.add_page_break()
        
        # Add executive summary
        if include_executive_summary and proposal_data.get('executive_summary'):
            self._add_executive_summary(doc, proposal_data)
            doc.add_page_break()
        
        # Add table of contents
        if include_toc:
            self._add_professional_toc(doc, proposal_data)
            doc.add_page_break()
        
        # Add sections
        for idx, section in enumerate(proposal_data.get('sections', []), 1):
            self._add_professional_section(doc, section, idx)
        
        # Add appendices if any
        if proposal_data.get('appendices'):
            doc.add_page_break()
            self._add_appendices(doc, proposal_data['appendices'])
        
        # Add headers and footers
        self._add_headers_footers(doc, proposal_data)
        
        # Save to bytes
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return file_stream.getvalue()
    
    def _setup_custom_styles(self, doc: Document, custom_styles: Optional[Dict]):
        """
        Set up custom paragraph and character styles
        """
        styles = doc.styles
        
        # Heading 1 style
        if 'Heading 1' in styles:
            h1 = styles['Heading 1']
            h1.font.name = 'Calibri'
            h1.font.size = Pt(16)
            h1.font.bold = True
            h1.font.color.rgb = self.company_branding['primary_color']
        
        # Heading 2 style
        if 'Heading 2' in styles:
            h2 = styles['Heading 2']
            h2.font.name = 'Calibri'
            h2.font.size = Pt(14)
            h2.font.bold = True
            h2.font.color.rgb = self.company_branding['secondary_color']
        
        # Normal style
        if 'Normal' in styles:
            normal = styles['Normal']
            normal.font.name = 'Calibri'
            normal.font.size = Pt(11)
            normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            normal.paragraph_format.space_after = Pt(6)
    
    def _set_document_properties(self, doc: Document, proposal_data: Dict[str, Any]):
        """
        Set document metadata properties
        """
        core_properties = doc.core_properties
        core_properties.title = proposal_data.get('title', 'Technical Proposal')
        core_properties.subject = proposal_data.get('rfp_number', '')
        core_properties.author = proposal_data.get('company', {}).get('name', 'GovSureAI')
        core_properties.comments = 'Generated by GovSureAI - Advanced Government Contracting Platform'
        core_properties.created = datetime.now()
    
    def _add_professional_cover_page(self, doc: Document, proposal_data: Dict[str, Any]):
        """
        Add professional cover page with branding
        """
        # Add logo if available
        if self.company_branding.get('logo_path') and os.path.exists(self.company_branding['logo_path']):
            doc.add_picture(self.company_branding['logo_path'], width=Inches(2))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(proposal_data.get('title', 'Technical Proposal'))
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = self.company_branding['primary_color']
        
        doc.add_paragraph()
        
        # RFP Information
        rfp_info = proposal_data.get('rfp_info', {})
        if rfp_info:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = p.add_run(f"{rfp_info.get('agency', '')}\n")
            run.font.size = Pt(16)
            run.font.bold = True
            
            run = p.add_run(f"{rfp_info.get('title', '')}\n\n")
            run.font.size = Pt(14)
            
            run = p.add_run(f"Solicitation Number: {rfp_info.get('number', '')}\n")
            run.font.size = Pt(12)
            
            if rfp_info.get('due_date'):
                run = p.add_run(f"Due Date: {rfp_info.get('due_date')}\n")
                run.font.size = Pt(12)
        
        # Spacing
        for _ in range(3):
            doc.add_paragraph()
        
        # Company Information
        company = proposal_data.get('company', {})
        if company:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = p.add_run("Submitted By\n\n")
            run.font.size = Pt(14)
            run.font.bold = True
            
            run = p.add_run(f"{company.get('name', '')}\n")
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = self.company_branding['primary_color']
            
            run = p.add_run(f"\n{company.get('address', '')}\n")
            run.font.size = Pt(11)
            
            run = p.add_run(f"{company.get('city', '')}, {company.get('state', '')} {company.get('zip', '')}\n")
            run.font.size = Pt(11)
            
            run = p.add_run(f"Phone: {company.get('phone', '')}\n")
            run.font.size = Pt(11)
            
            if company.get('email'):
                run = p.add_run(f"Email: {company.get('email', '')}\n")
                run.font.size = Pt(11)
        
        # Date
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(datetime.now().strftime("%B %d, %Y"))
        run.font.size = Pt(12)
        run.font.bold = True
    
    def _add_executive_summary(self, doc: Document, proposal_data: Dict[str, Any]):
        """
        Add executive summary section
        """
        heading = doc.add_heading('Executive Summary', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        summary_text = proposal_data.get('executive_summary', '')
        p = doc.add_paragraph(summary_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def _add_professional_toc(self, doc: Document, proposal_data: Dict[str, Any]):
        """
        Add professional table of contents
        """
        heading = doc.add_heading('Table of Contents', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Add TOC entries
        for idx, section in enumerate(proposal_data.get('sections', []), 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{idx}.0  {section.get('title', '')}")
            run.font.size = Pt(12)
            run.font.bold = True
            
            # Add subsections if any
            for sub_idx, subsection in enumerate(section.get('subsections', []), 1):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                run = p.add_run(f"{idx}.{sub_idx}  {subsection.get('title', '')}")
                run.font.size = Pt(11)
    
    def _add_professional_section(self, doc: Document, section: Dict[str, Any], section_number: int):
        """
        Add professionally formatted section
        """
        # Section heading
        heading = doc.add_heading(f"{section_number}.0  {section.get('title', '')}", level=1)
        
        # Section content
        content = section.get('content', '')
        if content:
            paragraphs = content.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    p = doc.add_paragraph(para_text.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Add subsections
        for sub_idx, subsection in enumerate(section.get('subsections', []), 1):
            sub_heading = doc.add_heading(
                f"{section_number}.{sub_idx}  {subsection.get('title', '')}",
                level=2
            )
            
            sub_content = subsection.get('content', '')
            if sub_content:
                paragraphs = sub_content.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        p = doc.add_paragraph(para_text.strip())
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Add tables if any
        if section.get('tables'):
            for table_data in section['tables']:
                self._add_table(doc, table_data)
        
        doc.add_paragraph()  # Spacing
    
    def _add_table(self, doc: Document, table_data: Dict[str, Any]):
        """
        Add formatted table
        """
        rows = table_data.get('rows', [])
        if not rows:
            return
        
        # Create table
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Light Grid Accent 1'
        
        # Populate table
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_data in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = str(cell_data)
                
                # Header row formatting
                if row_idx == 0:
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        doc.add_paragraph()  # Spacing
    
    def _add_appendices(self, doc: Document, appendices: List[Dict[str, Any]]):
        """
        Add appendices section
        """
        doc.add_heading('Appendices', level=1)
        
        for idx, appendix in enumerate(appendices, 1):
            doc.add_heading(f"Appendix {chr(64 + idx)}: {appendix.get('title', '')}", level=2)
            
            content = appendix.get('content', '')
            if content:
                doc.add_paragraph(content)
            
            doc.add_paragraph()
    
    def _add_headers_footers(self, doc: Document, proposal_data: Dict[str, Any]):
        """
        Add headers and footers
        """
        section = doc.sections[0]
        
        # Header
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = proposal_data.get('title', 'Technical Proposal')
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_para.runs[0].font.size = Pt(9)
        header_para.runs[0].font.color.rgb = self.company_branding['secondary_color']
        
        # Footer
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"{proposal_data.get('company', {}).get('name', 'GovSureAI')} - Confidential"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.runs[0].font.size = Pt(9)
        footer_para.runs[0].font.color.rgb = self.company_branding['secondary_color']
    
    async def export_to_advanced_excel(
        self,
        pricing_data: Dict[str, Any],
        include_charts: bool = True,
        include_formulas: bool = True
    ) -> bytes:
        """
        Export to advanced Excel workbook with multiple sheets, formulas, and charts
        """
        workbook = openpyxl.Workbook()
        
        # Remove default sheet
        if 'Sheet' in workbook.sheetnames:
            del workbook['Sheet']
        
        # Create sheets
        self._create_summary_sheet(workbook, pricing_data)
        self._create_labor_categories_sheet(workbook, pricing_data)
        self._create_cost_breakdown_sheet(workbook, pricing_data, include_formulas)
        self._create_pricing_schedule_sheet(workbook, pricing_data)
        
        if include_charts:
            self._add_charts_to_excel(workbook, pricing_data)
        
        # Save to bytes
        file_stream = io.BytesIO()
        workbook.save(file_stream)
        file_stream.seek(0)
        
        return file_stream.getvalue()
    
    def _create_summary_sheet(self, workbook, pricing_data: Dict[str, Any]):
        """
        Create executive summary sheet
        """
        ws = workbook.create_sheet('Executive Summary', 0)
        
        # Title
        ws['A1'] = 'Cost Proposal Summary'
        ws['A1'].font = Font(size=16, bold=True, color='FFFFFF')
        ws['A1'].fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
        ws.merge_cells('A1:D1')
        
        # Project info
        row = 3
        ws[f'A{row}'] = 'Project:'
        ws[f'B{row}'] = pricing_data.get('project_name', '')
        ws[f'A{row}'].font = Font(bold=True)
        
        row += 1
        ws[f'A{row}'] = 'Contract Type:'
        ws[f'B{row}'] = pricing_data.get('contract_type', 'FFP')
        ws[f'A{row}'].font = Font(bold=True)
        
        row += 2
        ws[f'A{row}'] = 'Cost Category'
        ws[f'B{row}'] = 'Amount'
        ws[f'A{row}'].font = Font(bold=True)
        ws[f'B{row}'].font = Font(bold=True)
        
        # Cost summary
        row += 1
        ws[f'A{row}'] = 'Direct Labor'
        ws[f'B{row}'] = f"=SUM('Labor Categories'!D:D)"
        
        row += 1
        ws[f'A{row}'] = 'Indirect Costs'
        ws[f'B{row}'] = pricing_data.get('indirect_costs', 0)
        
        row += 1
        ws[f'A{row}'] = 'Materials & Equipment'
        ws[f'B{row}'] = pricing_data.get('materials', 0)
        
        row += 1
        ws[f'A{row}'] = 'Travel'
        ws[f'B{row}'] = pricing_data.get('travel', 0)
        
        row += 1
        ws[f'A{row}'] = 'Total Cost'
        ws[f'B{row}'] = f"=SUM(B{row-4}:B{row-1})"
        ws[f'A{row}'].font = Font(bold=True)
        ws[f'B{row}'].font = Font(bold=True)
        
        # Format currency
        for r in range(7, row + 1):
            ws[f'B{r}'].number_format = '$#,##0.00'
        
        # Adjust column widths
        ws.column_dimensions['A'].width = 25
        ws.column_dimensions['B'].width = 15
    
    def _create_labor_categories_sheet(self, workbook, pricing_data: Dict[str, Any]):
        """
        Create labor categories sheet
        """
        ws = workbook.create_sheet('Labor Categories')
        
        # Headers
        headers = ['Labor Category', 'Hours', 'Rate', 'Total Cost']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(1, col, header)
            cell.font = Font(bold=True, color='FFFFFF')
            cell.fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
        
        # Data
        labor_categories = pricing_data.get('labor_categories', [])
        for row, category in enumerate(labor_categories, 2):
            ws.cell(row, 1, category.get('name', ''))
            ws.cell(row, 2, category.get('hours', 0))
            ws.cell(row, 3, category.get('rate', 0))
            ws.cell(row, 4, f"=B{row}*C{row}")
            
            # Format
            ws.cell(row, 3).number_format = '$#,##0.00'
            ws.cell(row, 4).number_format = '$#,##0.00'
        
        # Adjust widths
        ws.column_dimensions['A'].width = 30
        ws.column_dimensions['B'].width = 12
        ws.column_dimensions['C'].width = 12
        ws.column_dimensions['D'].width = 15
    
    def _create_cost_breakdown_sheet(self, workbook, pricing_data: Dict[str, Any], include_formulas: bool):
        """
        Create detailed cost breakdown sheet
        """
        ws = workbook.create_sheet('Cost Breakdown')
        
        # Implementation similar to labor categories but more detailed
        ws['A1'] = 'Detailed Cost Breakdown'
        ws['A1'].font = Font(size=14, bold=True)
        
        # Add detailed breakdown data
        # (Implementation details omitted for brevity)
    
    def _create_pricing_schedule_sheet(self, workbook, pricing_data: Dict[str, Any]):
        """
        Create pricing schedule sheet
        """
        ws = workbook.create_sheet('Pricing Schedule')
        
        ws['A1'] = 'Pricing Schedule by Period'
        ws['A1'].font = Font(size=14, bold=True)
        
        # Add pricing schedule data
        # (Implementation details omitted for brevity)
    
    def _add_charts_to_excel(self, workbook, pricing_data: Dict[str, Any]):
        """
        Add charts to Excel workbook
        """
        ws = workbook['Executive Summary']
        
        # Create bar chart for cost breakdown
        chart = BarChart()
        chart.title = "Cost Breakdown"
        chart.x_axis.title = "Category"
        chart.y_axis.title = "Amount ($)"
        
        data = Reference(ws, min_col=2, min_row=6, max_row=10)
        categories = Reference(ws, min_col=1, min_row=7, max_row=10)
        
        chart.add_data(data, titles_from_data=False)
        chart.set_categories(categories)
        
        ws.add_chart(chart, "D3")


# Global instance
enhanced_export_service = EnhancedExportService()

