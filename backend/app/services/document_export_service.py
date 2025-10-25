"""
Document Export Service - Export proposals to Word, Excel, PDF
"""
from typing import Dict, Any, List, Optional
from datetime import datetime
import io
import os

# Document generation libraries
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
except ImportError:
    pass


class DocumentExportService:
    """
    Export proposals and documents to various formats:
    - Word (.docx) - Formatted proposals with styles
    - Excel (.xlsx) - Cost breakdowns and pricing
    - PDF (.pdf) - Professional final output
    """
    
    def export_to_word(
        self,
        proposal_data: Dict[str, Any],
        include_cover_page: bool = True,
        include_toc: bool = True
    ) -> bytes:
        """
        Export proposal to Word document (.docx)
        
        Args:
            proposal_data: Proposal content with sections
            include_cover_page: Add professional cover page
            include_toc: Add table of contents
        
        Returns:
            bytes: Word document content
        """
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add cover page
        if include_cover_page:
            self._add_cover_page(doc, proposal_data)
            doc.add_page_break()
        
        # Add table of contents
        if include_toc:
            self._add_table_of_contents(doc, proposal_data)
            doc.add_page_break()
        
        # Add proposal sections
        for section in proposal_data.get('sections', []):
            self._add_section(doc, section)
        
        # Save to bytes
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return file_stream.getvalue()
    
    def _add_cover_page(self, doc: Document, proposal_data: Dict[str, Any]):
        """
        Add professional cover page
        """
        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(proposal_data.get('title', 'Technical Proposal'))
        run.font.size = Pt(24)
        run.font.bold = True
        
        doc.add_paragraph()  # Spacer
        
        # RFP Information
        rfp_info = proposal_data.get('rfp_info', {})
        if rfp_info:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"In Response to {rfp_info.get('agency', '')}\n")
            run.font.size = Pt(14)
            run = p.add_run(f"{rfp_info.get('title', '')}\n")
            run.font.size = Pt(14)
            run = p.add_run(f"RFP Number: {rfp_info.get('number', '')}\n")
            run.font.size = Pt(12)
        
        doc.add_paragraph()  # Spacer
        doc.add_paragraph()  # Spacer
        
        # Company Information
        company = proposal_data.get('company', {})
        if company:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Submitted By\n\n")
            run.font.size = Pt(14)
            run.font.bold = True
            
            run = p.add_run(f"{company.get('name', '')}\n")
            run.font.size = Pt(16)
            run.font.bold = True
            
            run = p.add_run(f"{company.get('address', '')}\n")
            run.font.size = Pt(12)
            run = p.add_run(f"{company.get('city', '')}, {company.get('state', '')} {company.get('zip', '')}\n")
            run.font.size = Pt(12)
            run = p.add_run(f"Phone: {company.get('phone', '')}\n")
            run.font.size = Pt(12)
        
        # Date
        doc.add_paragraph()  # Spacer
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(datetime.now().strftime("%B %d, %Y"))
        run.font.size = Pt(12)
    
    def _add_table_of_contents(self, doc: Document, proposal_data: Dict[str, Any]):
        """
        Add table of contents
        """
        heading = doc.add_heading('Table of Contents', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacer
        
        # Add section entries
        for i, section in enumerate(proposal_data.get('sections', []), 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}.0  {section.get('title', '')}")
            run.font.size = Pt(12)
            
            # Add subsections if any
            for j, subsection in enumerate(section.get('subsections', []), 1):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                run = p.add_run(f"{i}.{j}  {subsection.get('title', '')}")
                run.font.size = Pt(11)
    
    def _add_section(self, doc: Document, section: Dict[str, Any]):
        """
        Add a proposal section
        """
        # Section heading
        heading = doc.add_heading(section.get('title', ''), level=1)
        
        # Section content
        content = section.get('content', '')
        if content:
            # Split into paragraphs
            paragraphs = content.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    p = doc.add_paragraph(para_text.strip())
                    p.paragraph_format.space_after = Pt(12)
                    p.paragraph_format.line_spacing = 1.15
        
        # Add subsections
        for subsection in section.get('subsections', []):
            sub_heading = doc.add_heading(subsection.get('title', ''), level=2)
            
            sub_content = subsection.get('content', '')
            if sub_content:
                paragraphs = sub_content.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        p = doc.add_paragraph(para_text.strip())
                        p.paragraph_format.space_after = Pt(12)
                        p.paragraph_format.line_spacing = 1.15
        
        # Add tables if any
        for table_data in section.get('tables', []):
            self._add_table(doc, table_data)
        
        # Add page break after major sections
        if section.get('page_break', False):
            doc.add_page_break()
    
    def _add_table(self, doc: Document, table_data: Dict[str, Any]):
        """
        Add a table to the document
        """
        rows = table_data.get('rows', [])
        if not rows:
            return
        
        # Create table
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Light Grid Accent 1'
        
        # Fill table
        for i, row_data in enumerate(rows):
            for j, cell_value in enumerate(row_data):
                cell = table.rows[i].cells[j]
                cell.text = str(cell_value)
                
                # Bold header row
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
        
        doc.add_paragraph()  # Spacer after table
    
    def export_to_excel(
        self,
        pricing_data: Dict[str, Any],
        include_summary: bool = True
    ) -> bytes:
        """
        Export pricing/cost breakdown to Excel (.xlsx)
        
        Args:
            pricing_data: Pricing information with labor categories, rates, etc.
            include_summary: Add summary sheet
        
        Returns:
            bytes: Excel workbook content
        """
        workbook = openpyxl.Workbook()
        
        # Remove default sheet
        if 'Sheet' in workbook.sheetnames:
            del workbook['Sheet']
        
        # Add summary sheet
        if include_summary:
            self._add_summary_sheet(workbook, pricing_data)
        
        # Add labor categories sheet
        self._add_labor_categories_sheet(workbook, pricing_data)
        
        # Add cost breakdown sheet
        self._add_cost_breakdown_sheet(workbook, pricing_data)
        
        # Add pricing by year sheet
        if pricing_data.get('multi_year'):
            self._add_pricing_by_year_sheet(workbook, pricing_data)
        
        # Save to bytes
        file_stream = io.BytesIO()
        workbook.save(file_stream)
        file_stream.seek(0)
        
        return file_stream.getvalue()
    
    def _add_summary_sheet(self, workbook, pricing_data: Dict[str, Any]):
        """
        Add pricing summary sheet
        """
        sheet = workbook.create_sheet('Summary', 0)
        
        # Header
        sheet['A1'] = 'PRICING SUMMARY'
        sheet['A1'].font = Font(size=16, bold=True)
        sheet.merge_cells('A1:D1')
        
        # Proposal info
        row = 3
        sheet[f'A{row}'] = 'Proposal:'
        sheet[f'B{row}'] = pricing_data.get('proposal_title', '')
        sheet[f'A{row}'].font = Font(bold=True)
        
        row += 1
        sheet[f'A{row}'] = 'RFP Number:'
        sheet[f'B{row}'] = pricing_data.get('rfp_number', '')
        sheet[f'A{row}'].font = Font(bold=True)
        
        row += 1
        sheet[f'A{row}'] = 'Date:'
        sheet[f'B{row}'] = datetime.now().strftime("%B %d, %Y")
        sheet[f'A{row}'].font = Font(bold=True)
        
        row += 2
        
        # Pricing summary
        sheet[f'A{row}'] = 'TOTAL PRICING'
        sheet[f'A{row}'].font = Font(size=14, bold=True)
        row += 1
        
        # Add totals
        totals = pricing_data.get('totals', {})
        for label, value in totals.items():
            sheet[f'A{row}'] = label
            sheet[f'B{row}'] = value
            sheet[f'B{row}'].number_format = '$#,##0.00'
            row += 1
        
        # Style columns
        sheet.column_dimensions['A'].width = 25
        sheet.column_dimensions['B'].width = 20
    
    def _add_labor_categories_sheet(self, workbook, pricing_data: Dict[str, Any]):
        """
        Add labor categories and rates sheet
        """
        sheet = workbook.create_sheet('Labor Categories')
        
        # Header
        headers = ['Labor Category', 'Hourly Rate', 'Estimated Hours', 'Total Cost']
        for col, header in enumerate(headers, 1):
            cell = sheet.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True, color='FFFFFF')
            cell.fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
            cell.alignment = Alignment(horizontal='center')
        
        # Add labor categories
        labor_categories = pricing_data.get('labor_categories', [])
        for row, category in enumerate(labor_categories, 2):
            sheet.cell(row=row, column=1, value=category.get('name', ''))
            sheet.cell(row=row, column=2, value=category.get('hourly_rate', 0))
            sheet.cell(row=row, column=3, value=category.get('estimated_hours', 0))
            sheet.cell(row=row, column=4, value=category.get('total_cost', 0))
            
            # Format currency
            sheet.cell(row=row, column=2).number_format = '$#,##0.00'
            sheet.cell(row=row, column=4).number_format = '$#,##0.00'
        
        # Add totals row
        total_row = len(labor_categories) + 2
        sheet.cell(row=total_row, column=1, value='TOTAL')
        sheet.cell(row=total_row, column=1).font = Font(bold=True)
        sheet.cell(row=total_row, column=4, value=sum(cat.get('total_cost', 0) for cat in labor_categories))
        sheet.cell(row=total_row, column=4).number_format = '$#,##0.00'
        sheet.cell(row=total_row, column=4).font = Font(bold=True)
        
        # Style columns
        sheet.column_dimensions['A'].width = 30
        sheet.column_dimensions['B'].width = 15
        sheet.column_dimensions['C'].width = 18
        sheet.column_dimensions['D'].width = 15
    
    def _add_cost_breakdown_sheet(self, workbook, pricing_data: Dict[str, Any]):
        """
        Add detailed cost breakdown sheet
        """
        sheet = workbook.create_sheet('Cost Breakdown')
        
        # Header
        sheet['A1'] = 'DETAILED COST BREAKDOWN'
        sheet['A1'].font = Font(size=14, bold=True)
        sheet.merge_cells('A1:E1')
        
        # Column headers
        row = 3
        headers = ['Item', 'Description', 'Quantity', 'Unit Price', 'Total']
        for col, header in enumerate(headers, 1):
            cell = sheet.cell(row=row, column=col, value=header)
            cell.font = Font(bold=True)
        
        row += 1
        
        # Add cost items
        cost_items = pricing_data.get('cost_items', [])
        for item in cost_items:
            sheet.cell(row=row, column=1, value=item.get('name', ''))
            sheet.cell(row=row, column=2, value=item.get('description', ''))
            sheet.cell(row=row, column=3, value=item.get('quantity', 0))
            sheet.cell(row=row, column=4, value=item.get('unit_price', 0))
            sheet.cell(row=row, column=5, value=item.get('total', 0))
            
            # Format currency
            sheet.cell(row=row, column=4).number_format = '$#,##0.00'
            sheet.cell(row=row, column=5).number_format = '$#,##0.00'
            
            row += 1
        
        # Style columns
        sheet.column_dimensions['A'].width = 25
        sheet.column_dimensions['B'].width = 40
        sheet.column_dimensions['C'].width = 12
        sheet.column_dimensions['D'].width = 15
        sheet.column_dimensions['E'].width = 15
    
    def _add_pricing_by_year_sheet(self, workbook, pricing_data: Dict[str, Any]):
        """
        Add multi-year pricing sheet
        """
        sheet = workbook.create_sheet('Pricing by Year')
        
        # Header
        sheet['A1'] = 'MULTI-YEAR PRICING'
        sheet['A1'].font = Font(size=14, bold=True)
        
        # Years
        years = pricing_data.get('years', [])
        row = 3
        sheet.cell(row=row, column=1, value='Labor Category')
        sheet.cell(row=row, column=1).font = Font(bold=True)
        
        for col, year in enumerate(years, 2):
            sheet.cell(row=row, column=col, value=f"Year {year}")
            sheet.cell(row=row, column=col).font = Font(bold=True)
        
        # Add pricing by year
        row += 1
        labor_categories = pricing_data.get('labor_categories', [])
        for category in labor_categories:
            sheet.cell(row=row, column=1, value=category.get('name', ''))
            
            for col, year in enumerate(years, 2):
                price = category.get(f'year_{year}_cost', 0)
                sheet.cell(row=row, column=col, value=price)
                sheet.cell(row=row, column=col).number_format = '$#,##0.00'
            
            row += 1
    
    def export_to_pdf(
        self,
        proposal_data: Dict[str, Any],
        include_cover_page: bool = True
    ) -> bytes:
        """
        Export proposal to PDF (.pdf)
        
        Args:
            proposal_data: Proposal content
            include_cover_page: Add cover page
        
        Returns:
            bytes: PDF document content
        """
        file_stream = io.BytesIO()
        doc = SimpleDocTemplate(file_stream, pagesize=letter)
        
        # Container for PDF elements
        elements = []
        
        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=30,
            alignment=1  # Center
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=12,
            spaceBefore=12
        )
        
        # Add cover page
        if include_cover_page:
            elements.append(Paragraph(proposal_data.get('title', 'Technical Proposal'), title_style))
            elements.append(Spacer(1, 0.5*inch))
            
            rfp_info = proposal_data.get('rfp_info', {})
            if rfp_info:
                elements.append(Paragraph(f"In Response to {rfp_info.get('agency', '')}", styles['Normal']))
                elements.append(Paragraph(rfp_info.get('title', ''), styles['Normal']))
                elements.append(Paragraph(f"RFP Number: {rfp_info.get('number', '')}", styles['Normal']))
            
            elements.append(PageBreak())
        
        # Add sections
        for section in proposal_data.get('sections', []):
            elements.append(Paragraph(section.get('title', ''), heading_style))
            
            content = section.get('content', '')
            if content:
                paragraphs = content.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        elements.append(Paragraph(para_text.strip(), styles['Normal']))
                        elements.append(Spacer(1, 0.1*inch))
            
            elements.append(Spacer(1, 0.2*inch))
        
        # Build PDF
        doc.build(elements)
        file_stream.seek(0)
        
        return file_stream.getvalue()


# Global instance
document_export_service = DocumentExportService()

